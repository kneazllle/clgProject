from flask import Flask, request, jsonify, send_file
import requests
from flask_mysqldb import MySQL
import MySQLdb
from flask_cors import CORS
from docx import Document
import mammoth

app = Flask(__name__)

# Configure CORS to allow requests from React frontend
cors = CORS(app, resources={r"/api/*": {"origins": "http://localhost:3000"}})

db = MySQLdb.Connect(host="containers-us-west-78.railway.app", port=5480,
                     user="root", passwd="F09DY9R7wJEsodY9LB1B", db="railway")

# Get all the services


@app.route('/api/services', methods=["GET"])
def services():
    cur = db.cursor()
    cur.execute('SELECT * FROM services')
    row_headers = [x[0] for x in cur.description]
    rv = cur.fetchall()
    json_data = []
    for result in rv:
        json_data.append(dict(zip(row_headers, result)))
    cur.close()
    return jsonify(json_data)

# Get forms of a particular service


@app.route('/api/forms', methods=["GET"])
def get_forms():
    # Send json object {"service_id": "..."}
    Service = request.args.get('service_id')
    print(type(Service))
    print(Service)
    cur = db.cursor()
    cur.execute(
        "SELECT services.service_id, services.service_name, forms.form_id, forms.form_name, forms.form_link FROM services INNER JOIN forms ON services.service_id = forms.service_id WHERE forms.service_id = %s;", [Service])
    row_headers = [x[0] for x in cur.description]
    rv = cur.fetchall()
    json_data = []
    for result in rv:
        json_data.append(dict(zip(row_headers, result)))
    cur.close()
    return jsonify(json_data)

# Get all queries for a form


@app.route('/api/form-details', methods=["GET"])
def get_form_details():
    # Send json object {"form_id":"..."}
    form_id = request.args.get('form_id')
    print(form_id)
    cur = db.cursor()
    cur.execute("SELECT * FROM forms WHERE form_id = %s;", [form_id])
    row_headers = [x[0] for x in cur.description]
    rv = cur.fetchall()
    json_data = []
    for result in rv:
        json_data.append(dict(zip(row_headers, result)))
    cur.execute(
        "SELECT * FROM input_ques WHERE ques_id IN (SELECT form_query_id FROM form_queries WHERE form_id = %s);", [form_id])
    row_headers = [x[0] for x in cur.description]
    rv = cur.fetchall()
    for result in rv:
        json_data.append(dict(zip(row_headers, result)))
    cur.close()
    return jsonify(json_data)


# Return the contents of final doc
@app.route('/api/final-content', methods=["POST"])
def final_content():
    form_details = request.json                         # Under Progress
    form_id = form_details["form_id"]
    print(type(form_details))
    cur = db.cursor()
    cur.execute("SELECT form_link FROM forms where form_id = %s;", [form_id])
    row_headers = [x[0] for x in cur.description]
    rv = cur.fetchall()
    json_data = []
    for result in rv:
        json_data.append(dict(zip(row_headers, result)))
    cur.close()
    print(json_data[0]["form_link"])
    response = requests.get(json_data[0]["form_link"])
    with open('../docs/localfile.docx', 'wb') as f:
        f.write(response.content)
    doc = Document('../docs/localfile.docx')
    test = list(form_details.keys())
    
    test.sort(reverse=True)
    
    for key in test:
        old = "#"+str(key)
        new = str(form_details[key])
        # print(old, new)
        for p in doc.paragraphs:
            if old in p.text:
                # print(old)
                inline = p.runs
                for i in range(len(inline)):
                    if old in inline[i].text:
                        # print(old)
                        res = inline[i].text.replace(old, new)
                        inline[i].text = res
    doc.save("docs/Output2.docx")
    
    f = open('docs/Output2.docx', 'rb')
    
    docx_content = mammoth.convert_to_html(f)
    print(docx_content.value)
    # docx_content.close()

    # fullText = []
    # for para in doc.paragraphs:
    #     fullText.append(para.text)
    # fullText = '\n'.join(fullText)
    # print(fullText)
    return jsonify({'content': docx_content.value})


# Return the final doc


@app.route('/api/final-form', methods=["POST"])
def final_form():
    contents = request.get_json()
    print(contents)
    with open('docs/Output2.docx', 'w') as file:
        file.write(contents)
    # doc = Document('docs/localfile.docx')
    # for key, value in form_details.items():
    #     for paragraph in doc.paragraphs:
    #         paragraph.text = paragraph.text.replace(
    #             "#"+str(key)+'#', str(value))

    # doc.save("docs/Output2.docx")
    return send_file('docs\Output2.docx', as_attachment=True)


if __name__ == '__main__':
    app.run(debug=True)
